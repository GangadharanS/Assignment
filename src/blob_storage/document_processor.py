"""
Document processing module for extracting text from various file formats.
"""
import io
from pathlib import Path
from typing import Optional

# Try importing optional dependencies
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocumentProcessor:
    """Processes various document formats and extracts text content."""
    
    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".doc", ".md", ".json", ".csv"}
    
    def __init__(self):
        """Initialize the document processor."""
        pass
    
    def extract_text(self, file_path: str = None, file_content: bytes = None, 
                     filename: str = None) -> str:
        """
        Extract text from a document.
        
        Args:
            file_path: Path to the file (optional if file_content provided)
            file_content: Raw bytes of the file (optional if file_path provided)
            filename: Original filename (needed when using file_content to determine type)
            
        Returns:
            Extracted text content
        """
        if file_path:
            path = Path(file_path)
            extension = path.suffix.lower()
            with open(file_path, "rb") as f:
                content = f.read()
        elif file_content and filename:
            extension = Path(filename).suffix.lower()
            content = file_content
        else:
            raise ValueError("Either file_path or (file_content and filename) must be provided")
        
        # Route to appropriate extractor
        if extension == ".pdf":
            return self._extract_pdf(content)
        elif extension in {".docx", ".doc"}:
            return self._extract_docx(content)
        elif extension in {".txt", ".md", ".json", ".csv"}:
            return self._extract_text_file(content)
        else:
            # Try to decode as text
            return self._extract_text_file(content)
    
    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF content."""
        if not HAS_PYPDF:
            raise ImportError(
                "pypdf is required for PDF processing. "
                "Install with: pip install pypdf"
            )
        
        text_parts = []
        pdf_file = io.BytesIO(content)
        reader = pypdf.PdfReader(pdf_file)
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx"
            )
        
        doc_file = io.BytesIO(content)
        doc = DocxDocument(doc_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_text_file(self, content: bytes) -> str:
        """Extract text from plain text files."""
        # Try different encodings
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Fallback: decode with errors ignored
        return content.decode("utf-8", errors="ignore")
    
    def is_supported(self, filename: str) -> bool:
        """Check if a file type is supported."""
        extension = Path(filename).suffix.lower()
        return extension in self.SUPPORTED_EXTENSIONS or extension in {".txt", ".md"}


# Singleton instance
document_processor = DocumentProcessor()


